from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# 标题
title = doc.add_heading('Claude Code Game Studios - 开发流程', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('游戏：从奴隶到星际霸主（暂定名）', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('引擎：Godot 4 | 视角：2D俯视角 | 目标：一个月出Demo', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# ---- 流程总览 ----
doc.add_heading('流程总览', level=1)
doc.add_paragraph('以下流程按顺序执行，每步完成后检查产出物再进入下一步。')

steps = [
    ('☑ 已完成', '填写游戏想法收集表', 'design/game-idea.md'),
    ('□ 第1步', '/brainstorm — 头脑风暴', '把想法发展成完整游戏概念文档'),
    ('□ 第2步', '/setup-engine — 配置引擎', '指定 Godot 4.6，初始化项目结构'),
    ('□ 第3步', '/map-systems — 拆解系统', '列出所有系统模块（发兵、占点、AI、UI…）'),
    ('□ 第4步', '/create-epics — 创建史诗', '系统分批，先做哪些后做哪些'),
    ('□ 第5步', '/create-stories — 创建故事', '拆成具体可执行的开发任务'),
    ('□ 第6步', '/sprint-plan — Sprint 计划', '排期，定第一个Sprint的内容'),
    ('□ 第7步', '/dev-story — 开始开发', '按Story一个个实现，写出战争系统核心'),
]

table = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
hdr = table.rows[0].cells
hdr[0].text = '状态'
hdr[1].text = '命令'
hdr[2].text = '说明'
hdr[0].paragraphs[0].runs[0].bold = True
hdr[1].paragraphs[0].runs[0].bold = True
hdr[2].paragraphs[0].runs[0].bold = True

for status, cmd, desc in steps:
    row = table.add_row().cells
    row[0].text = status
    row[1].text = cmd
    row[2].text = desc

# 调整列宽
for row in table.rows:
    row.cells[0].width = Cm(2.5)
    row.cells[1].width = Cm(5.5)
    row.cells[2].width = Cm(9)

doc.add_paragraph('')

# ---- 详细说明 ----
doc.add_heading('每步详细说明', level=1)

doc.add_heading('第1步：/brainstorm', level=2)
p = doc.add_paragraph()
p.add_run('输入：').bold = True
p.add_run('design/game-idea.md（已填好的游戏想法）')
p = doc.add_paragraph()
p.add_run('产出：').bold = True
p.add_run('完整的游戏概念文档（design/gdd/game-concept.md）')
p = doc.add_paragraph()
p.add_run('说明：').bold = True
p.add_run('Game-Studios 会用专业框架（MDA、玩家心理学等）把你的想法扩展成完整概念。你会被问一些问题，回答就行。')

doc.add_heading('第2步：/setup-engine', level=2)
p = doc.add_paragraph()
p.add_run('输入：').bold = True
p.add_run('选择 Godot 4')
p = doc.add_paragraph()
p.add_run('产出：').bold = True
p.add_run('引擎版本锁定、项目目录结构初始化、技术偏好配置')
p = doc.add_paragraph()
p.add_run('说明：').bold = True
p.add_run('指定 Godot 4.6，自动创建 src/ assets/ tests/ 等目录，配置引擎文档参考。你的电脑上已有 Godot 4.6.2。')

doc.add_heading('第3步：/map-systems', level=2)
p = doc.add_paragraph()
p.add_run('输入：').bold = True
p.add_run('游戏概念文档')
p = doc.add_paragraph()
p.add_run('产出：').bold = True
p.add_run('系统清单 + 依赖关系图')
p = doc.add_paragraph()
p.add_run('说明：').bold = True
p.add_run('把你的战争系统拆成子模块——拖线发兵、占点产兵、AI敌人、胜负判定、UI显示等。')

doc.add_heading('第4步：/create-epics', level=2)
p = doc.add_paragraph()
p.add_run('输入：').bold = True
p.add_run('系统清单')
p = doc.add_paragraph()
p.add_run('产出：').bold = True
p.add_run('Epic 列表（按优先级排列）')
p = doc.add_paragraph()
p.add_run('说明：').bold = True
p.add_run('把系统分组，决定哪些 Demo 必须做、哪些以后再说。')

doc.add_heading('第5步：/create-stories', level=2)
p = doc.add_paragraph()
p.add_run('输入：').bold = True
p.add_run('每个 Epic')
p = doc.add_paragraph()
p.add_run('产出：').bold = True
p.add_run('可执行的 Story 文件（含验收标准）')
p = doc.add_paragraph()
p.add_run('说明：').bold = True
p.add_run('就是把"做发兵系统"拆成"实现拖动检测"、"实现兵力派遣"、"实现星球占领逻辑"这种颗粒度。')

doc.add_heading('第6步：/sprint-plan', level=2)
p = doc.add_paragraph()
p.add_run('输入：').bold = True
p.add_run('Story 列表')
p = doc.add_paragraph()
p.add_run('产出：').bold = True
p.add_run('第一个 Sprint 计划')
p = doc.add_paragraph()
p.add_run('说明：').bold = True
p.add_run('排好第一周做什么、第二周做什么。')

doc.add_heading('第7步：/dev-story', level=2)
p = doc.add_paragraph()
p.add_run('输入：').bold = True
p.add_run('Story 文件')
p = doc.add_paragraph()
p.add_run('产出：').bold = True
p.add_run('实际代码 + 测试')
p = doc.add_paragraph()
p.add_run('说明：').bold = True
p.add_run('这是真正写代码的步骤。Game-Studios 会自动分派对应的程序员代理干活。一个 Story 做完再做下一个。')

doc.add_paragraph('')

# ---- 快捷参考 ----
doc.add_heading('Demo目标提醒', level=2)
items = [
    '只做战争系统核心：拖线发兵 + 占点产兵 + 全歼制胜',
    '不做第一阶段的 RPG 部分',
    '用临时美术素材跑通玩法就行',
    '单机、无限时、像素风',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')

# ---- 注意事项 ----
doc.add_heading('注意事项', level=2)
notes = [
    '每一步都在 Claude Code 终端里输入对应命令（如 /brainstorm）执行',
    '必须在 d:\\AIcode\\Claude-Code-Game-Studios 目录下打开 Claude Code',
    '每步完成后检查产出文件是否存在，再进下一步',
    '中间觉得不对随时叫小C调整',
    'Game-Studios 的代理们会在需要时自动出现，不用手动调用',
]
for note in notes:
    doc.add_paragraph(note, style='List Number')

# 保存
desktop = os.path.expanduser('~/Desktop')
filepath = os.path.join(desktop, 'Game-开发流程.docx')
doc.save(filepath)
print(f'已保存到: {filepath}')
